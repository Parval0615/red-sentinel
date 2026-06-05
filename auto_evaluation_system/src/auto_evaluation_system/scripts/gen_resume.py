from auto_evaluation_system.bootstrap import setup_paths  # noqa: F401
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.size = Pt(10.5)
style.font.name = '微软雅黑'
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.2

ACCENT = RGBColor(0x1C, 0x49, 0x6F)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x88, 0x88, 0x88)
LIGHT_BG = RGBColor(0xF0, 0xF3, 0xF7)

# ====== HEADER (2-col: info + photo) ======
header_table = doc.add_table(rows=1, cols=2)
header_table.autofit = True

# Left: all text info
h_left = header_table.cell(0, 0)
h_left.width = Cm(13)

name_p = h_left.paragraphs[0]
name_p.paragraph_format.space_after = Pt(3)
r = name_p.add_run('陈中泽')
r.font.size = Pt(26)
r.font.bold = True
r.font.name = '微软雅黑'
r.font.color.rgb = ACCENT

contact_p = h_left.add_paragraph()
contact_p.paragraph_format.space_after = Pt(1)
r = contact_p.add_run('18477424221  |  1440746764@qq.com')
r.font.size = Pt(9.5)
r.font.name = '微软雅黑'
r.font.color.rgb = GRAY

school_p = h_left.add_paragraph()
school_p.paragraph_format.space_after = Pt(1)
r = school_p.add_run('杭州电子科技大学 网络空间安全学院 密码学（研一）  |  福州大学 信息管理与信息系统（本科）')
r.font.size = Pt(9.5)
r.font.name = '微软雅黑'
r.font.color.rgb = GRAY

intent_p = h_left.add_paragraph()
intent_p.paragraph_format.space_after = Pt(0)
r1 = intent_p.add_run('意向岗位：')
r1.font.size = Pt(9.5)
r1.font.name = '微软雅黑'
r1.font.color.rgb = GRAY
r2 = intent_p.add_run('AI安全工程师（实习）/ 大模型安全')
r2.font.size = Pt(9.5)
r2.font.bold = True
r2.font.name = '微软雅黑'
r2.font.color.rgb = ACCENT

# Right: photo placeholder
h_right = header_table.cell(0, 1)
h_right.width = Cm(3)

p_photo = h_right.paragraphs[0]
p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_photo.paragraph_format.space_before = Pt(2)
r_empty = p_photo.add_run(' ')
r_empty.font.size = Pt(2)

# Vertical divider between left and right (optional, clean)
# Just set a minimum height on the right cell
tcPr = h_right._element.get_or_add_tcPr()
# right-aligned cell content

p_label = h_right.add_paragraph()
p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_label.paragraph_format.space_after = Pt(1)
r_label = p_label.add_run('[ 照片 ]')
r_label.font.size = Pt(12)
r_label.font.name = '微软雅黑'
r_label.font.color.rgb = GRAY

p_size = h_right.add_paragraph()
p_size.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_size = p_size.add_run('一寸')
r_size.font.size = Pt(8)
r_size.font.name = '微软雅黑'
r_size.font.color.rgb = GRAY

# Add border around photo area
tcPr = h_right._element.get_or_add_tcPr()
# Light background
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'F5F6FA')
shd.set(qn('w:val'), 'clear')
tcPr.append(shd)
# Borders (only around the right cell, not between cells)
tcBorders = OxmlElement('w:tcBorders')
for edge in ['top', 'left', 'bottom', 'right']:
    border = OxmlElement(f'w:{edge}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '6')
    border.set(qn('w:space'), '0')
    border.set(qn('w:color'), '1C496F')
    tcBorders.append(border)
tcPr.append(tcBorders)

# Remove outer borders from header table
tbl = header_table._element
tblPr = tbl.find(qn('w:tblPr'))
if tblPr is None:
    tblPr = OxmlElement('w:tblPr')
    tbl.insert(0, tblPr)
tblBorders = OxmlElement('w:tblBorders')
for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    border = OxmlElement(f'w:{edge}')
    border.set(qn('w:val'), 'nil')
    tblBorders.append(border)
tblPr.append(tblBorders)

# Divider
div = doc.add_paragraph()
div.paragraph_format.space_after = Pt(2)
div.paragraph_format.space_before = Pt(0)
pPr = div._element.get_or_add_pPr()
pBdr = pPr.makeelement(qn('w:pBdr'), {})
bottom = pBdr.makeelement(qn('w:bottom'), {
    qn('w:val'): 'single', qn('w:sz'): '8',
    qn('w:space'): '1', qn('w:color'): '1C496F',
})
pBdr.append(bottom)
pPr.append(pBdr)

# ====== Helpers ======
def section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.name = '微软雅黑'
    r.font.color.rgb = ACCENT

def proj_title(name, date):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(name)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.name = '微软雅黑'
    r.font.color.rgb = DARK
    # date on same line if room, or on next
    r2 = p.add_run('    ' + date)
    r2.font.size = Pt(9)
    r2.font.name = '微软雅黑'
    r2.font.color.rgb = GRAY

def bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run('- ' + text)
    r.font.size = Pt(9.5)
    r.font.name = '微软雅黑'

# ====== PROJECTS ======
section_title('项目经历')

# --- Project 1: AI Security ---
proj_title('AI Agent Runtime Security -- 企业RAG Agent安全防护与评测系统', '2026.02 - 至今')

bullet('从0构建面向企业RAG Agent的AI安全防护系统，覆盖Prompt注入、RAG文档投毒、'
       '工具滥用、输出幻觉、多租户泄露五大攻击面，实现纵深防御链')

bullet('设计三层智能防火墙：L1规则引擎（复合词对+正则，<1ms）+ L2 LLM语义分类 '
       '+ L3操作粒度后检，105条攻击载荷拦截率93.3%，0%误拦；复现Indirect Prompt '
       'Injection完整攻击链（HTML/Email/Markdown/PDF四格式25场景，RAG检索-LLM污染-'
       '危险工具调用-Policy Engine拦截），5/5策略测试通过')

bullet('实现审计日志SHA256哈希链防篡改 + Ed25519数字签名完整性验证，精确检测篡改条目索引；'
       '工具供应链Ed25519签名验证（4/4工具+6/6篡改检测），校验失败拒绝加载并写入审计日志')

bullet('技术栈：Python, LangGraph(StateGraph+Checkpointer), ChromaDB(BM25+Vector RRF融合'
       '+BGE-Reranker), Streamlit, HuggingFace, EasyOCR')

# --- Project 2: Bearing Fault Diagnosis ---
proj_title('基于迁移学习的高速列车轴承故障诊断模型建模', '2024.09 - 2024.12  |  "华为杯"研究生数学建模 国家二等奖')

bullet('针对高速列车真实故障数据稀缺问题，构建跨工况智能故障诊断系统。从160+源域信号'
       '提取时/频/时频域多维特征，搭建XGBoost多任务分类模型；自研类条件CORAL+伪标签'
       '自训练迁移学习策略，实现目标域无标签数据精准分类')

bullet('设计"小波降噪-重采样-滑动分窗"三级预处理链路统一多源采样基准，提取时域、频域'
       '（频谱散度/滚降）、时频域特征；对比逻辑回归/SVM/XGBoost等5种模型，网格搜索+'
       '分层K折交叉验证调优+SMOTE平衡，最优模型故障类型/尺寸分类准确率达95.4%/95.0%')

bullet('提出类条件CORAL对齐+伪标签自训练方法克服源/目标域分布差异，对16个无标签轴承'
       '信号完成标注；通过SHAP值、t-SNE可视化及协方差差距分析验证模型决策的物理可解释性')

bullet('技术栈：Python(NumPy/SciPy/scikit-learn/XGBoost), 信号处理(小波降噪/FFT/STFT), '
       '迁移学习(CORAL/伪标签自训练), SHAP/t-SNE可解释性分析')

# --- Project 3: Math Modeling ---
proj_title('基于LSTM和随机森林的奥运奖牌分布预测', '2025.02  |  美国大学生数学建模竞赛 H奖（国际二等奖）')

bullet('运用LSTM+蒙特卡罗随机失活法预测下届奥运会奖牌分布及首次获奖国家，随机森林'
       '识别关键体育项目对奖牌分布的影响权重，负责Python数据处理、模型调参与可视化')

# ====== SKILLS + EDUCATION (2-column, clean) ======
section_title('专业技能与教育背景')

info_table = doc.add_table(rows=1, cols=2)
info_table.autofit = True

# --- Left: Skills ---
left = info_table.cell(0, 0)
left.width = Cm(9.5)

items_left = [
    ('编程语言：', 'Python, C, SQL'),
    ('安全攻防：', 'Prompt注入 / RAG投毒 / Indirect Injection / 工具滥用 / 审计日志防篡改'),
    ('密码学：', 'Ed25519数字签名 / SHA256哈希链 / 完整性验证'),
    ('框架工具：', 'LangGraph, LangChain, ChromaDB, Streamlit'),
    ('机器学习：', 'scikit-learn, XGBoost, SHAP, 信号处理(FFT/STFT/小波降噪)'),
    ('英语：', 'CET-6'),
]

first = True
for label, value in items_left:
    p = left.paragraphs[0] if first else left.add_paragraph()
    first = False
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(label)
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.name = '微软雅黑'
    r2 = p.add_run(value)
    r2.font.size = Pt(9.5)
    r2.font.name = '微软雅黑'

# --- Right: Education + Honors ---
right = info_table.cell(0, 1)
right.width = Cm(7.5)

items_right = [
    ('研', '杭州电子科技大学 网络空间安全学院 密码学'),
    ('', '2025.09 - 2028.06（预计）'),
    ('本', '福州大学 信息管理与信息系统'),
    ('', '2021.09 - 2025.06'),
    ('', ''),
    ('', '"华为杯"研究生数学建模竞赛 国家二等奖'),
    ('', '美国大学生数学建模竞赛 H奖（国际二等）'),
    ('', '全国大学生数学竞赛 省级一等奖'),
]

first = True
for label, value in items_right:
    p = right.paragraphs[0] if first else right.add_paragraph()
    first = False
    if not value and not label:
        p.paragraph_format.space_after = Pt(2)
        continue
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    if label:
        r = p.add_run(label + ': ')
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.name = '微软雅黑'
    r2 = p.add_run(value)
    r2.font.size = Pt(9.5)
    r2.font.name = '微软雅黑'

# Hide table borders
tbl = info_table._element
tblPr = tbl.find(qn('w:tblPr'))
if tblPr is None:
    tblPr = OxmlElement('w:tblPr')
    tbl.insert(0, tblPr)
tblBorders = OxmlElement('w:tblBorders')
for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    border = OxmlElement(f'w:{edge}')
    border.set(qn('w:val'), 'nil')
    tblBorders.append(border)
tblPr.append(tblBorders)

output_path = str(__import__('pathlib').Path(__file__).resolve().parents[2] / 'reports' / '陈中泽简历.docx')
doc.save(output_path)
print('Saved OK, size:', os.path.getsize(output_path), 'bytes')
